import io
import os
import json
from typing import List, Dict, Any, Optional

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None


class LoadedDocument:
    """Represents an extracted document with metadata and content sections."""
    def __init__(self, filename: str, full_text: str, sections: List[Dict[str, Any]], metadata: Optional[Dict[str, Any]] = None):
        self.filename = filename
        self.full_text = full_text
        self.sections = sections  # List of {"label": str (e.g. "Page 1" or "Sheet 1"), "text": str}
        self.metadata = metadata or {}

    def to_dict(self) -> Dict[str, Any]:
        return {
            "filename": self.filename,
            "full_text": self.full_text,
            "sections": self.sections,
            "metadata": self.metadata,
        }


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> LoadedDocument:
    """Extract text page-by-page from a PDF file."""
    if not PdfReader:
        raise ImportError("pypdf is required to read PDF files.")
    
    reader = PdfReader(io.BytesIO(file_bytes))
    sections = []
    full_text_parts = []
    
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        page_text = page_text.strip()
        if page_text:
            sections.append({
                "label": f"Page {i + 1}",
                "text": page_text,
                "page_number": i + 1
            })
            full_text_parts.append(f"--- [Page {i + 1}] ---\n{page_text}")
            
    full_text = "\n\n".join(full_text_parts)
    return LoadedDocument(
        filename=filename,
        full_text=full_text,
        sections=sections,
        metadata={"total_pages": len(reader.pages), "extracted_pages": len(sections), "type": "PDF"}
    )


def extract_text_from_docx(file_bytes: bytes, filename: str) -> LoadedDocument:
    """Extract paragraphs and tables from a Word (.docx) document."""
    if not docx:
        raise ImportError("python-docx is required to read .docx files.")
    
    doc = docx.Document(io.BytesIO(file_bytes))
    sections = []
    full_text_parts = []
    
    # Paragraphs
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if para_texts:
        para_content = "\n".join(para_texts)
        sections.append({"label": "Document Body", "text": para_content})
        full_text_parts.append(para_content)
        
    # Tables
    for t_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(row_data))
        table_text = "\n".join(table_rows)
        if table_text.strip():
            sections.append({"label": f"Table {t_idx + 1}", "text": table_text})
            full_text_parts.append(f"--- [Table {t_idx + 1}] ---\n{table_text}")
            
    full_text = "\n\n".join(full_text_parts)
    return LoadedDocument(
        filename=filename,
        full_text=full_text,
        sections=sections,
        metadata={"paragraphs_count": len(para_texts), "tables_count": len(doc.tables), "type": "Word"}
    )


def extract_text_from_spreadsheet(file_bytes: bytes, filename: str, ext: str) -> LoadedDocument:
    """Extract data from Excel (.xlsx, .xls) or CSV files into formatted tables."""
    sections = []
    full_text_parts = []
    
    if ext == ".csv":
        try:
            df = pd.read_csv(io.BytesIO(file_bytes))
            csv_md = df.to_markdown(index=False) if hasattr(df, "to_markdown") else df.to_string(index=False)
            sections.append({"label": "CSV Data", "text": csv_md})
            full_text_parts.append(csv_md)
        except Exception:
            text = file_bytes.decode("utf-8", errors="replace")
            sections.append({"label": "CSV Raw", "text": text})
            full_text_parts.append(text)
    else:
        # Excel
        try:
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheet_md = df.to_markdown(index=False) if hasattr(df, "to_markdown") else df.to_string(index=False)
                label = f"Sheet: {sheet_name}"
                sections.append({"label": label, "text": sheet_md})
                full_text_parts.append(f"--- [{label}] ---\n{sheet_md}")
        except Exception as e:
            text = f"Error reading Excel sheet: {str(e)}"
            sections.append({"label": "Error", "text": text})
            full_text_parts.append(text)
            
    full_text = "\n\n".join(full_text_parts)
    return LoadedDocument(
        filename=filename,
        full_text=full_text,
        sections=sections,
        metadata={"type": "Spreadsheet", "sections_count": len(sections)}
    )


def extract_text_from_generic_text(file_bytes: bytes, filename: str, ext: str) -> LoadedDocument:
    """Extract text from plain text, code, markdown, json, or any generic extension."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "utf-16"]
    text = None
    for enc in encodings:
        try:
            text = file_bytes.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
            
    if text is None:
        text = file_bytes.decode("utf-8", errors="ignore")
        
    # If JSON, try to format nicely
    if ext == ".json":
        try:
            parsed = json.loads(text)
            text = json.dumps(parsed, indent=2)
        except Exception:
            pass
            
    lines = text.splitlines()
    sections = []
    
    # Group into reasonable chunks/sections if large
    if len(lines) > 200:
        chunk_size = 150
        for i in range(0, len(lines), chunk_size):
            chunk_lines = lines[i : i + chunk_size]
            sec_text = "\n".join(chunk_lines)
            sections.append({
                "label": f"Lines {i + 1}-{min(i + chunk_size, len(lines))}",
                "text": sec_text
            })
    else:
        sections.append({"label": "Full Content", "text": text})
        
    return LoadedDocument(
        filename=filename,
        full_text=text,
        sections=sections,
        metadata={"line_count": len(lines), "type": "Text/Code/Data"}
    )


def load_file_content(file_bytes: bytes, filename: str) -> LoadedDocument:
    """Universal dispatcher to load and parse any file type."""
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes, filename)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes, filename)
    elif ext in [".xlsx", ".xls", ".csv"]:
        return extract_text_from_spreadsheet(file_bytes, filename, ext)
    else:
        # Plain text, Markdown, Python, JSON, YAML, HTML, CSS, Log, Config, or any arbitrary file extension
        return extract_text_from_generic_text(file_bytes, filename, ext)
